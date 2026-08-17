# -*- coding: utf-8 -*-
"""按催化剂公司公文处理规范(v2.0个人拟稿版)生成测试公文
只含主体部分：标题、主送机关、正文、附件说明、署名成文日期、附件
红头版头和版记由公文系统自动套格式，不制作"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLACK = RGBColor(0, 0, 0)
doc = Document()

# ============ 页面设置：A4，上37/下35/左28/右26mm ============
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin, sec.bottom_margin = Mm(37), Mm(35)
sec.left_margin, sec.right_margin = Mm(28), Mm(26)


def set_run(run, font, size_pt):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = False
    run.font.color.rgb = BLACK
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)


def para(text, font='仿宋_GB2312', size=16, align=WD_ALIGN_PARAGRAPH.LEFT,
         indent_chars=0, line=27):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(line)          # 固定值行距
    pf.space_before = pf.space_after = Pt(0)
    if indent_chars:
        ind = p._element.get_or_add_pPr().get_or_add_ind()
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))
    if text:
        r = p.add_run(text)
        set_run(r, font, size)
    return p


# ============ 标题：2号方正小标宋，居中，不加粗，多行梯形 ============
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing = Pt(33)
r_t1 = p_title.add_run('关于开展2026年度物资需求计划')
set_run(r_t1, '方正小标宋简体', 22)
r_t1.add_break()
r_t2 = p_title.add_run('编报工作的通知')
set_run(r_t2, '方正小标宋简体', 22)

# ============ 主送机关：标题下空一行，居左顶格，全角冒号 ============
para('', size=16, line=27)
para('机关各部门、各分（子）公司：', size=16)

# ============ 正文：3号仿宋，行距固定27磅，首行缩进2字符 ============
para('为深入贯彻落实集团公司物资采购管理有关要求，全面提升物资需求计划编报质量，'
     '保障生产经营物资供应，现就2026年度物资需求计划编报工作有关事项通知如下。',
     size=16, indent_chars=2)

# 第一层：黑体
para('一、总体要求', font='黑体', size=16, indent_chars=2)
para('坚持“以需定采、计划先行”原则，各单位要严格执行《中国石化物资采购管理规定》'
     '（中国石化制〔2019〕46号）有关要求，做到需求准确、依据充分、编报及时，'
     '切实提高计划的严肃性和准确性。', size=16, indent_chars=2)

para('二、编报范围及要求', font='黑体', size=16, indent_chars=2)
# 第二层：楷体
para('（一）编报范围', font='楷体', size=16, indent_chars=2)
para('各单位2026年度生产经营、检维修、技措技改及安全生产所需的各类物资，'
     '均须纳入年度需求计划编报范围（见附件1）。', size=16, indent_chars=2)
para('（二）编报要求', font='楷体', size=16, indent_chars=2)
# 第三层：仿宋
para('1.需求计划须以年度生产经营计划、检修计划为依据，逐项列明物资名称、规格型号、'
     '数量、需用时间及技术要求。', size=16, indent_chars=2)
para('2.单项物资估算金额达到限额标准的，须同步编制采购方式建议及投资、成本费用'
     '预算落实情况说明（见附件2）。', size=16, indent_chars=2)

para('三、时间安排', font='黑体', size=16, indent_chars=2)
para('各单位请于2026年9月10日前完成本单位需求计划审核汇总，'
     '经负责人签字确认后报送物装中心计划管理部。', size=16, indent_chars=2)

para('四、有关要求', font='黑体', size=16, indent_chars=2)
para('各单位要高度重视，加强组织领导，明确专人负责，确保编报工作按期高质量完成。',
     size=16, indent_chars=2)

# ============ 附件说明：正文下空一行，左空二字 ============
para('', size=16, line=27)
para('附件：1.2026年度物资需求计划表', size=16, indent_chars=2)
para('2.物资需求计划编报说明', size=16, indent_chars=2)

# ============ 署名、成文日期：右空二字，日期首字右移二字，同页 ============
para('', size=16, line=27)
p_sign = para('中国石化催化剂有限公司　　　　', size=16, align=WD_ALIGN_PARAGRAPH.RIGHT)
p_sign2 = para('物资装备中心　　　　', size=16, align=WD_ALIGN_PARAGRAPH.RIGHT)
p_sign2.paragraph_format.keep_with_next = True
p_sign.paragraph_format.keep_with_next = True
para('2026年8月17日　　', size=16, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ============ 附件另面编排 ============
doc.add_page_break()
para('附件1', font='黑体', size=16)
para('', size=16, line=27)
para('2026年度物资需求计划表', font='方正小标宋简体', size=22,
     align=WD_ALIGN_PARAGRAPH.CENTER, line=33)
para('', size=16, line=27)

tbl = doc.add_table(rows=4, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['序号', '物资名称', '规格型号', '数量', '需用时间']
for j, h in enumerate(headers):
    cell = tbl.rows[0].cells[j]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(h)
    set_run(r, '黑体', 12)
rows_data = [['1', '示例物资A', 'XXX-100', '10', '2026年3月'],
             ['2', '示例物资B', 'XXX-200', '5', '2026年6月'],
             ['3', '示例物资C', 'XXX-300', '20', '2026年9月']]
for i, row in enumerate(rows_data, 1):
    for j, v in enumerate(row):
        cell = tbl.rows[i].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(v)
        set_run(r, '仿宋_GB2312', 12)

out = r'C:/Users/LENOVO/.qwenworkcn/workspace/mswn5u7alzmrdykm/outputs/公文排版测试_物资需求计划编报通知_v5.docx'
doc.save(out)
print('已生成:', out)
