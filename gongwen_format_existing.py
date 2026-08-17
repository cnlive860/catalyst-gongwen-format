# -*- coding: utf-8 -*-
"""按catalyst-gongwen-format技能格式化绿色采购授课稿件"""
import re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r'D:/Desktop/工作/3.绿色采购管理/绿色采购培训/绿色采购专篇授课稿件.docx'
BLACK = RGBColor(0, 0, 0)

d = Document(SRC)

# ============ 版面：A4，上37/下35/左28/右26mm ============
sec = d.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin, sec.bottom_margin = Mm(37), Mm(35)
sec.left_margin, sec.right_margin = Mm(28), Mm(26)


def set_run(r, font, size):
    r.font.name = 'Times New Roman' if font == '宋体' else font
    r.font.size = Pt(size)
    r.font.bold = False
    r.font.color.rgb = BLACK
    r.font.italic = False
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)


def set_para(p, font, size, align=WD_ALIGN_PARAGRAPH.LEFT, indent=2, line=27):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(line)
    pf.space_before = pf.space_after = Pt(0)
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    if indent:
        ind = pPr.makeelement(qn('w:ind'), {})
        ind.set(qn('w:firstLineChars'), str(indent * 100))
        pPr.append(ind)
    for r in p.runs:
        set_run(r, font, size)
    # 段内可能有超链接等非run文本节点，统一字体到此为止


re_l1 = re.compile(r'^(开场白|板块[一二三四五六]|结语)')
re_l2 = re.compile(r'^第\s*\d+\s*页')

for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        # 空段保持为空，行距规范化
        p.paragraph_format.line_spacing = Pt(27)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
        continue
    if t == '绿色企业创建之绿色采购专篇':
        set_para(p, '方正小标宋简体', 22, WD_ALIGN_PARAGRAPH.CENTER, indent=0, line=33)
    elif re_l1.match(t):
        set_para(p, '黑体', 16)          # 第一层：黑体
    elif re_l2.match(t):
        set_para(p, '楷体', 16)          # 第二层：楷体
    else:
        set_para(p, '仿宋_GB2312', 16)   # 正文：3号仿宋

# ============ 文末两行落款 + 成文日期 ============
# 若文末最后一段非空则先补一空行
if d.paragraphs[-1].text.strip():
    d.add_paragraph('')

p1 = d.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p1.paragraph_format.line_spacing = Pt(27)
r = p1.add_run('中国石化催化剂有限公司　　　　')
set_run(r, '仿宋_GB2312', 16)
p1.paragraph_format.keep_with_next = True

p2 = d.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.line_spacing = Pt(27)
p2.paragraph_format.keep_with_next = True
r = p2.add_run('物资装备中心　　　　')
set_run(r, '仿宋_GB2312', 16)

p3 = d.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p3.paragraph_format.line_spacing = Pt(27)
r = p3.add_run('2026年8月17日　　')
set_run(r, '仿宋_GB2312', 16)

d.save(SRC)
print('已按规范格式化并保存:', SRC)
